#!/usr/bin/env python3
"""Generate the ArtBot Hardware Plan Word document with diagrams."""

import os
import sys
import tempfile

# Ensure gpy venv
sys.path.insert(0, "/Users/evantobias/.python-envs/globalenv/lib/python3.13/site-packages")

import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
import matplotlib.patches as mpatches
from matplotlib.patches import FancyBboxPatch, FancyArrowPatch, Rectangle, Circle
import numpy as np

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT

OUTDIR = os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))), "docs")
IMGDIR = tempfile.mkdtemp(prefix="artbot_diagrams_")


# ═══════════════════════════════════════════════════════════════
# DIAGRAM GENERATORS
# ═══════════════════════════════════════════════════════════════

def draw_system_overview():
    """Main machine layout diagram."""
    fig, ax = plt.subplots(figsize=(8, 6))
    ax.set_xlim(-0.5, 10.5)
    ax.set_ylim(-1.5, 8.5)
    ax.set_aspect("equal")
    ax.axis("off")
    ax.set_title("ArtBot System Layout — Front View", fontsize=14, fontweight="bold", pad=15)

    # Whiteboard frame
    board = Rectangle((0, 0), 10, 7, linewidth=2.5, edgecolor="#333", facecolor="#f5f5f0", zorder=1)
    ax.add_patch(board)
    ax.text(5, -0.4, "180 cm", ha="center", fontsize=9, color="#666")
    ax.text(-0.4, 3.5, "120 cm", ha="center", fontsize=9, color="#666", rotation=90)

    # Buffer zone
    buf = Rectangle((0.28, 0.28), 9.44, 6.44, linewidth=1, edgecolor="#4CAF50",
                     facecolor="none", linestyle="--", zorder=2, alpha=0.7)
    ax.add_patch(buf)
    ax.text(5, 0.5, "Safe Zone: 170 × 110 cm (5cm buffer)", ha="center", fontsize=7, color="#4CAF50")

    # Y-axis rails (left and right)
    ax.plot([0.5, 0.5], [0.5, 6.5], color="#1565C0", linewidth=4, solid_capstyle="round", zorder=3)
    ax.plot([9.5, 9.5], [0.5, 6.5], color="#1565C0", linewidth=4, solid_capstyle="round", zorder=3)
    ax.text(0.5, 6.7, "Y-Rail\n(2020)", ha="center", fontsize=6, color="#1565C0")
    ax.text(9.5, 6.7, "Y-Rail\n(2020)", ha="center", fontsize=6, color="#1565C0")

    # Y motors at top
    motor_y_left = FancyBboxPatch((0.1, 6.8), 0.8, 0.6, boxstyle="round,pad=0.05",
                                   facecolor="#E53935", edgecolor="#B71C1C", linewidth=1.5, zorder=5)
    motor_y_right = FancyBboxPatch((9.1, 6.8), 0.8, 0.6, boxstyle="round,pad=0.05",
                                    facecolor="#E53935", edgecolor="#B71C1C", linewidth=1.5, zorder=5)
    ax.add_patch(motor_y_left)
    ax.add_patch(motor_y_right)
    ax.text(0.5, 7.1, "Motor\nY-Left", ha="center", fontsize=5.5, color="white", fontweight="bold", zorder=6)
    ax.text(9.5, 7.1, "Motor\nY-Right", ha="center", fontsize=5.5, color="white", fontweight="bold", zorder=6)

    # Y-axis belts (dashed lines from motors to rail)
    rail_y = 4.0
    ax.plot([0.5, 0.5], [6.8, rail_y + 0.15], color="#E53935", linewidth=1.5, linestyle="--", zorder=2)
    ax.plot([9.5, 9.5], [6.8, rail_y + 0.15], color="#E53935", linewidth=1.5, linestyle="--", zorder=2)
    ax.text(0.15, 5.5, "Belt", fontsize=6, color="#E53935", rotation=90)
    ax.text(9.85, 5.5, "Belt", fontsize=6, color="#E53935", rotation=90)

    # X-axis rail (horizontal)
    ax.plot([0.5, 9.5], [rail_y, rail_y], color="#FF8F00", linewidth=6, solid_capstyle="round", zorder=4)
    ax.text(5, rail_y + 0.3, "X-Rail (2040 V-Slot, 180 cm)", ha="center", fontsize=7, color="#FF8F00", fontweight="bold")

    # X motor on left end of rail
    motor_x = FancyBboxPatch((0.8, rail_y - 0.45), 0.7, 0.5, boxstyle="round,pad=0.05",
                              facecolor="#7B1FA2", edgecolor="#4A148C", linewidth=1.5, zorder=5)
    ax.add_patch(motor_x)
    ax.text(1.15, rail_y - 0.2, "Motor\nX", ha="center", fontsize=5.5, color="white", fontweight="bold", zorder=6)

    # Pen carriage
    carriage_x = 5.5
    carriage = FancyBboxPatch((carriage_x - 0.5, rail_y - 0.7), 1.0, 0.6,
                               boxstyle="round,pad=0.05", facecolor="#0288D1",
                               edgecolor="#01579B", linewidth=1.5, zorder=5)
    ax.add_patch(carriage)
    ax.text(carriage_x, rail_y - 0.4, "Pen\nCarriage", ha="center", fontsize=5.5,
            color="white", fontweight="bold", zorder=6)

    # Pen tip
    ax.plot([carriage_x, carriage_x], [rail_y - 0.7, rail_y - 1.1], color="#01579B",
            linewidth=2, zorder=5)
    ax.plot(carriage_x, rail_y - 1.15, "v", color="#01579B", markersize=8, zorder=5)
    ax.text(carriage_x + 0.3, rail_y - 1.0, "Pen", fontsize=6, color="#01579B")

    # Direct drive note
    ax.text(10.2, 2.0, "No counter-\nweights\n(NEMA 23\ndirect drive)", fontsize=6, color="#2E7D32", ha="left")

    # Arrows showing motion
    ax.annotate("", xy=(7.5, rail_y - 0.4), xytext=(3.5, rail_y - 0.4),
                arrowprops=dict(arrowstyle="<->", color="#7B1FA2", lw=1.5))
    ax.text(5.5, rail_y - 0.55, "X motion", ha="center", fontsize=6, color="#7B1FA2")

    ax.annotate("", xy=(0.1, 5.5), xytext=(0.1, 2.5),
                arrowprops=dict(arrowstyle="<->", color="#E53935", lw=1.5))
    ax.text(-0.15, 4.0, "Y motion", ha="center", fontsize=6, color="#E53935", rotation=90)

    # Legend
    legend_items = [
        mpatches.Patch(color="#E53935", label="Y-Motors (×2)"),
        mpatches.Patch(color="#7B1FA2", label="X-Motor (×1)"),
        mpatches.Patch(color="#0288D1", label="Pen Carriage + Servo"),
        mpatches.Patch(color="#FF8F00", label="X-Rail (2040)"),
        mpatches.Patch(color="#1565C0", label="Y-Rails (2020)"),
    ]
    ax.legend(handles=legend_items, loc="lower right", fontsize=6, framealpha=0.9)

    plt.tight_layout()
    path = os.path.join(IMGDIR, "system_overview.png")
    fig.savefig(path, dpi=200, bbox_inches="tight")
    plt.close(fig)
    return path


def draw_architecture_comparison():
    """Side-by-side comparison of the 4 architectures."""
    fig, axes = plt.subplots(1, 4, figsize=(11, 3.5))
    fig.suptitle("Motion Architecture Comparison", fontsize=13, fontweight="bold", y=1.02)

    for ax in axes:
        ax.set_xlim(0, 10)
        ax.set_ylim(0, 8)
        ax.set_aspect("equal")
        ax.axis("off")

    # --- A: Polargraph ---
    ax = axes[0]
    ax.set_title("A: Polargraph\n(REJECTED)", fontsize=8, fontweight="bold", color="#C62828")
    ax.add_patch(Rectangle((1, 0.5), 8, 6, edgecolor="#ccc", facecolor="#fafafa", lw=1))
    ax.plot(1.5, 7, "s", color="#E53935", markersize=10)
    ax.plot(8.5, 7, "s", color="#E53935", markersize=10)
    ax.text(1.5, 7.5, "M1", ha="center", fontsize=6, fontweight="bold")
    ax.text(8.5, 7.5, "M2", ha="center", fontsize=6, fontweight="bold")
    ax.plot([1.5, 5], [7, 3], color="#E53935", lw=1, linestyle="--")
    ax.plot([8.5, 5], [7, 3], color="#E53935", lw=1, linestyle="--")
    ax.plot(5, 3, "o", color="#0288D1", markersize=10, zorder=5)
    ax.text(5, 2.2, "Gondola", ha="center", fontsize=6)
    # Show bad accuracy zone
    ax.add_patch(Rectangle((1, 0.5), 8, 2, facecolor="#FFCDD2", alpha=0.5))
    ax.text(5, 1.2, "Low accuracy\nzone", ha="center", fontsize=5, color="#C62828")

    # --- B: CoreXY ---
    ax = axes[1]
    ax.set_title("B: CoreXY\n(REJECTED)", fontsize=8, fontweight="bold", color="#C62828")
    ax.add_patch(Rectangle((1, 0.5), 8, 6, edgecolor="#ccc", facecolor="#fafafa", lw=1))
    ax.plot(1.5, 7, "s", color="#E53935", markersize=10)
    ax.plot(8.5, 0.8, "s", color="#E53935", markersize=10)
    ax.text(1.5, 7.5, "M1", ha="center", fontsize=6, fontweight="bold")
    ax.text(8.5, 0.2, "M2", ha="center", fontsize=6, fontweight="bold")
    # Belt path (simplified)
    belt_x = [1.5, 8.5, 8.5, 5, 5, 1.5, 1.5, 5, 5, 8.5]
    belt_y = [7, 7, 0.8, 0.8, 4, 4, 7, 7, 4, 4]
    ax.plot(belt_x[:5], belt_y[:5], color="#FF8F00", lw=0.8, linestyle="--", alpha=0.7)
    ax.plot(5, 4, "o", color="#0288D1", markersize=10, zorder=5)
    ax.text(5, 3.2, "Carriage", ha="center", fontsize=6)
    ax.text(5, 1.5, "~12m belt path\nStretch problem", ha="center", fontsize=5, color="#C62828")

    # --- C: Single-Y Cartesian ---
    ax = axes[2]
    ax.set_title("C: Single-Y\n(REJECTED)", fontsize=8, fontweight="bold", color="#C62828")
    ax.add_patch(Rectangle((1, 0.5), 8, 6, edgecolor="#ccc", facecolor="#fafafa", lw=1))
    ax.plot([1.5, 1.5], [1, 6.5], color="#1565C0", lw=3)
    ax.plot(1.5, 7, "s", color="#E53935", markersize=10)
    ax.text(1.5, 7.5, "M-Y", ha="center", fontsize=6, fontweight="bold")
    ax.plot([1.5, 8.5], [4, 3.5], color="#FF8F00", lw=3)  # Tilted rail!
    ax.plot(5, 3.75, "o", color="#0288D1", markersize=10, zorder=5)
    ax.text(5, 1.5, "Rail tilts!\n(racking)", ha="center", fontsize=5, color="#C62828")

    # --- D: Dual-Y Cartesian ---
    ax = axes[3]
    ax.set_title("D: Dual-Y Cartesian\n(CHOSEN)", fontsize=8, fontweight="bold", color="#2E7D32")
    ax.add_patch(Rectangle((1, 0.5), 8, 6, edgecolor="#2E7D32", facecolor="#E8F5E9", lw=2))
    ax.plot([1.5, 1.5], [1, 6.5], color="#1565C0", lw=3)
    ax.plot([8.5, 8.5], [1, 6.5], color="#1565C0", lw=3)
    ax.plot(1.5, 7, "s", color="#E53935", markersize=10)
    ax.plot(8.5, 7, "s", color="#E53935", markersize=10)
    ax.text(1.5, 7.5, "M-Y1", ha="center", fontsize=6, fontweight="bold")
    ax.text(8.5, 7.5, "M-Y2", ha="center", fontsize=6, fontweight="bold")
    ax.plot([1.5, 8.5], [4, 4], color="#FF8F00", lw=3)  # Perfectly level!
    ax.plot(5, 4, "o", color="#0288D1", markersize=10, zorder=5)
    ax.text(5, 4.5, "Level rail", ha="center", fontsize=5, color="#2E7D32")
    ax.text(5, 1.5, "Precise\nScalable", ha="center", fontsize=5, color="#2E7D32", fontweight="bold")

    plt.tight_layout()
    path = os.path.join(IMGDIR, "architecture_comparison.png")
    fig.savefig(path, dpi=200, bbox_inches="tight")
    plt.close(fig)
    return path


def draw_direct_drive_diagram():
    """Side view showing the direct-drive Y-axis (no counterweights)."""
    fig, ax = plt.subplots(figsize=(6, 5))
    ax.set_xlim(-1, 11)
    ax.set_ylim(-2, 9)
    ax.set_aspect("equal")
    ax.axis("off")
    ax.set_title("Direct-Drive Y-Axis — Side View (No Counterweights)", fontsize=11, fontweight="bold", pad=10)

    # Board (side view - thin rectangle)
    ax.add_patch(Rectangle((2, 0), 0.3, 7, facecolor="#E0E0E0", edgecolor="#333", lw=1.5))
    ax.text(2.15, -0.5, "Board\n(vertical)", ha="center", fontsize=7)

    # Y-Rail (on front of board)
    ax.plot([2.3, 2.3], [0.5, 6.5], color="#1565C0", lw=4, zorder=3)
    ax.text(3.0, 6.5, "Y-Rail", fontsize=7, color="#1565C0")

    # NEMA 23 Motor at top (bigger than NEMA 17)
    motor = FancyBboxPatch((1.6, 7.0), 1.4, 0.9, boxstyle="round,pad=0.05",
                            facecolor="#E53935", edgecolor="#B71C1C", lw=1.5, zorder=5)
    ax.add_patch(motor)
    ax.text(2.3, 7.45, "NEMA 23\nMotor Y", ha="center", fontsize=6, color="white", fontweight="bold", zorder=6)

    # X-Rail (cross section - small square on the rail)
    rail_y = 4.0
    rail = FancyBboxPatch((2.1, rail_y - 0.3), 0.8, 0.6, boxstyle="round,pad=0.02",
                           facecolor="#FF8F00", edgecolor="#E65100", lw=1.5, zorder=4)
    ax.add_patch(rail)
    ax.text(3.3, rail_y, "X-Rail\n(end view)", fontsize=7, color="#FF8F00")

    # Belt from motor to rail
    ax.plot([2.3, 2.3], [7.0, rail_y + 0.3], color="#E53935", lw=1.5, linestyle="--", zorder=2)
    ax.text(1.3, 5.5, "Belt", fontsize=6, color="#E53935")

    # Idler at bottom
    ax.plot(2.3, 0.5, "o", color="#333", markersize=8, zorder=5)
    ax.text(1.3, 0.3, "Idler", fontsize=6)

    # Belt from rail to idler
    ax.plot([2.3, 2.3], [rail_y - 0.3, 0.5], color="#E53935", lw=1.5, linestyle="--", zorder=2)

    # Gravity arrow
    ax.annotate("", xy=(2.5, rail_y - 1.2), xytext=(2.5, rail_y - 0.3),
                arrowprops=dict(arrowstyle="->", color="#C62828", lw=2))
    ax.text(3.0, rail_y - 1.0, "Gravity", fontsize=6, color="#C62828")

    # Motor torque arrow (upward)
    ax.annotate("", xy=(2.0, rail_y + 0.8), xytext=(2.0, rail_y + 0.3),
                arrowprops=dict(arrowstyle="->", color="#2E7D32", lw=2))
    ax.text(0.5, rail_y + 0.6, "Motor\ntorque", fontsize=6, color="#2E7D32")

    # Info text
    ax.text(5.5, 6.0, "Direct Drive (No Counterweights):", fontsize=9, fontweight="bold")
    ax.text(5.5, 5.3, "NEMA 23 motors (100–300 Ncm)", fontsize=8)
    ax.text(5.5, 4.6, "provide enough torque to drive", fontsize=8)
    ax.text(5.5, 3.9, "the Y-axis directly against gravity.", fontsize=8)
    ax.text(5.5, 2.8, "Simpler build:", fontsize=8, fontweight="bold", color="#2E7D32")
    ax.text(5.5, 2.2, "• No pulleys, cables, or weights", fontsize=8, color="#2E7D32")
    ax.text(5.5, 1.6, "• No weight calibration needed", fontsize=8, color="#2E7D32")
    ax.text(5.5, 0.6, "Note: Rail drops on power loss.", fontsize=8, style="italic", color="#C62828")
    ax.text(5.5, 0.0, "Acceptable for pen plotter.", fontsize=8, style="italic", color="#C62828")

    plt.tight_layout()
    path = os.path.join(IMGDIR, "direct_drive.png")
    fig.savefig(path, dpi=200, bbox_inches="tight")
    plt.close(fig)
    return path


def draw_pen_mechanism():
    """Pen lift mechanism diagram."""
    fig, axes = plt.subplots(1, 2, figsize=(8, 3.5))
    fig.suptitle("Pen Lift Mechanism (Z-Axis)", fontsize=12, fontweight="bold", y=1.02)

    for ax in axes:
        ax.set_xlim(0, 10)
        ax.set_ylim(0, 7)
        ax.set_aspect("equal")
        ax.axis("off")

    # --- Pen DOWN ---
    ax = axes[0]
    ax.set_title("Pen DOWN (Drawing)", fontsize=9, color="#2E7D32", fontweight="bold")

    # Board surface
    ax.add_patch(Rectangle((0, 0), 0.5, 7, facecolor="#E0E0E0", edgecolor="#333"))
    ax.text(0.25, 0.3, "Board", rotation=90, ha="center", fontsize=6)

    # Pen holder bracket
    ax.add_patch(Rectangle((2, 2), 5, 3, facecolor="#BBDEFB", edgecolor="#1565C0", lw=1.5))
    ax.text(4.5, 4.2, "Pen Carriage", ha="center", fontsize=7)

    # Pen body
    ax.add_patch(Rectangle((1.0, 3.0), 3.0, 1.0, facecolor="#FFD54F", edgecolor="#F57F17", lw=1.5))
    ax.text(2.5, 3.5, "Pen", ha="center", fontsize=7, fontweight="bold")

    # Pen tip touching board
    ax.plot([0.5, 1.0], [3.5, 3.5], color="#F57F17", lw=2)
    ax.plot(0.5, 3.5, ">", color="#F57F17", markersize=10, zorder=5)

    # Spring (zigzag)
    spring_x = [4.0, 4.3, 4.0, 4.3, 4.0, 4.3, 4.0]
    spring_y = [3.2, 3.3, 3.4, 3.5, 3.6, 3.7, 3.8]
    ax.plot(spring_x, spring_y, color="#2E7D32", lw=2)
    ax.text(4.6, 3.5, "Spring\n(pushes pen\ntoward board)", fontsize=5, color="#2E7D32")

    # Servo (released)
    ax.add_patch(FancyBboxPatch((5.5, 2.5), 1.2, 1.5, boxstyle="round,pad=0.05",
                                 facecolor="#CE93D8", edgecolor="#7B1FA2", lw=1.5))
    ax.text(6.1, 3.25, "Servo\n(released)", ha="center", fontsize=5, color="#4A148C")

    # Arrow showing spring force
    ax.annotate("", xy=(0.8, 2.5), xytext=(2.5, 2.5),
                arrowprops=dict(arrowstyle="->", color="#2E7D32", lw=2))
    ax.text(1.5, 2.1, "Spring force →", fontsize=6, color="#2E7D32")

    # --- Pen UP ---
    ax = axes[1]
    ax.set_title("Pen UP (Traveling)", fontsize=9, color="#C62828", fontweight="bold")

    # Board surface
    ax.add_patch(Rectangle((0, 0), 0.5, 7, facecolor="#E0E0E0", edgecolor="#333"))
    ax.text(0.25, 0.3, "Board", rotation=90, ha="center", fontsize=6)

    # Pen holder bracket
    ax.add_patch(Rectangle((2, 2), 5, 3, facecolor="#BBDEFB", edgecolor="#1565C0", lw=1.5))
    ax.text(4.5, 4.2, "Pen Carriage", ha="center", fontsize=7)

    # Pen body (retracted - gap from board)
    ax.add_patch(Rectangle((2.0, 3.0), 3.0, 1.0, facecolor="#FFD54F", edgecolor="#F57F17", lw=1.5))
    ax.text(3.5, 3.5, "Pen", ha="center", fontsize=7, fontweight="bold")

    # Gap between pen and board
    ax.annotate("", xy=(0.6, 3.5), xytext=(1.9, 3.5),
                arrowprops=dict(arrowstyle="<->", color="#C62828", lw=1.5))
    ax.text(1.0, 3.0, "Gap", fontsize=6, color="#C62828")

    # Spring (compressed)
    spring_x = [5.0, 5.2, 5.0, 5.2, 5.0]
    spring_y = [3.3, 3.4, 3.5, 3.6, 3.7]
    ax.plot(spring_x, spring_y, color="#2E7D32", lw=2)

    # Servo (engaged, pulling pen back)
    ax.add_patch(FancyBboxPatch((5.5, 2.5), 1.2, 1.5, boxstyle="round,pad=0.05",
                                 facecolor="#7B1FA2", edgecolor="#4A148C", lw=1.5))
    ax.text(6.1, 3.25, "Servo\n(engaged)", ha="center", fontsize=5, color="white", fontweight="bold")

    # Arrow showing servo pull
    ax.annotate("", xy=(5.5, 3.5), xytext=(5.0, 3.5),
                arrowprops=dict(arrowstyle="<-", color="#7B1FA2", lw=2))
    ax.text(5.5, 2.1, "← Servo retracts pen", fontsize=6, color="#7B1FA2")

    plt.tight_layout()
    path = os.path.join(IMGDIR, "pen_mechanism.png")
    fig.savefig(path, dpi=200, bbox_inches="tight")
    plt.close(fig)
    return path


def draw_electronics_diagram():
    """Electronics wiring block diagram."""
    fig, ax = plt.subplots(figsize=(9, 6))
    ax.set_xlim(0, 14)
    ax.set_ylim(0, 10)
    ax.set_aspect("equal")
    ax.axis("off")
    ax.set_title("Electronics & Wiring Block Diagram", fontsize=13, fontweight="bold", pad=10)

    def box(ax, x, y, w, h, label, color, textcolor="white", fontsize=7):
        b = FancyBboxPatch((x, y), w, h, boxstyle="round,pad=0.1",
                            facecolor=color, edgecolor="black", lw=1.5, zorder=3)
        ax.add_patch(b)
        ax.text(x + w/2, y + h/2, label, ha="center", va="center",
                fontsize=fontsize, color=textcolor, fontweight="bold", zorder=4)

    # PSU
    box(ax, 0.5, 4.0, 2.0, 1.5, "24V 15A\nPower\nSupply", "#37474F")

    # Arduino + CNC Shield (center)
    box(ax, 4.0, 3.5, 3.0, 2.5, "Arduino Uno\n+\nCNC Shield V3\n(GRBL)", "#1565C0", fontsize=8)

    # PC/Laptop
    box(ax, 4.5, 7.5, 2.0, 1.2, "PC / Laptop\n(G-Code Sender)", "#333")
    ax.annotate("", xy=(5.5, 7.5), xytext=(5.5, 6.1),
                arrowprops=dict(arrowstyle="<->", color="#333", lw=1.5))
    ax.text(6.0, 6.8, "USB", fontsize=7, color="#333")

    # Power line
    ax.annotate("", xy=(4.0, 4.75), xytext=(2.5, 4.75),
                arrowprops=dict(arrowstyle="->", color="#C62828", lw=2))
    ax.text(3.0, 5.1, "24V", fontsize=7, color="#C62828", fontweight="bold")

    # Drivers
    driver_x = 8.5
    colors = ["#E53935", "#E53935", "#7B1FA2"]
    labels = ["TB6600\n(Y-Driver)", "TB6600\n(A=Y Clone)", "TB6600\n(X-Driver)"]
    for i, (col, lbl) in enumerate(zip(colors, labels)):
        y = 6.5 - i * 2.2
        box(ax, driver_x, y, 1.8, 1.2, lbl, col, fontsize=6)
        ax.annotate("", xy=(driver_x, y + 0.6), xytext=(7.0, 4.75),
                    arrowprops=dict(arrowstyle="->", color="#555", lw=1, connectionstyle="arc3,rad=0.1"))

    # Motors
    motor_x = 11.5
    motor_labels = ["Motor\nY-Left", "Motor\nY-Right\n(COIL\nREVERSED)", "Motor\nX"]
    motor_colors = ["#B71C1C", "#B71C1C", "#4A148C"]
    for i, (lbl, col) in enumerate(zip(motor_labels, motor_colors)):
        y = 6.5 - i * 2.2
        box(ax, motor_x, y, 1.8, 1.2, lbl, col, fontsize=5.5)
        ax.annotate("", xy=(motor_x, y + 0.6), xytext=(driver_x + 1.8, y + 0.6),
                    arrowprops=dict(arrowstyle="->", color=col, lw=1.5))

    # Servo
    box(ax, 8.5, 0.3, 1.8, 1.0, "SG90 Servo\n(Pen Z)", "#00695C", fontsize=6)
    ax.annotate("", xy=(8.5, 0.8), xytext=(7.0, 3.5),
                arrowprops=dict(arrowstyle="->", color="#00695C", lw=1,
                                connectionstyle="arc3,rad=-0.2"))
    ax.text(7.5, 1.8, "PWM", fontsize=6, color="#00695C")

    # Limit switches
    box(ax, 1.0, 0.5, 2.0, 1.2, "Limit Switches\nX-min, Y-min\n(Microswitches)", "#FF8F00", textcolor="black", fontsize=6)
    ax.annotate("", xy=(3.0, 1.1), xytext=(4.0, 3.5),
                arrowprops=dict(arrowstyle="->", color="#FF8F00", lw=1,
                                connectionstyle="arc3,rad=0.2"))

    # Warning box
    box(ax, 0.5, 8.0, 3.5, 1.2, "⚠ Y-Right motor:\nReverse ONE coil pair\n(swap A+ and A−)", "#FFF3E0", textcolor="#BF360C", fontsize=6)

    plt.tight_layout()
    path = os.path.join(IMGDIR, "electronics.png")
    fig.savefig(path, dpi=200, bbox_inches="tight")
    plt.close(fig)
    return path


def draw_assembly_phases():
    """Assembly timeline / phases diagram."""
    fig, ax = plt.subplots(figsize=(10, 4))
    ax.set_xlim(0, 12)
    ax.set_ylim(0, 3)
    ax.axis("off")
    ax.set_title("Assembly Phases", fontsize=13, fontweight="bold", pad=10)

    phases = [
        ("Phase 1\nFrame", "Mount rails\nSquare corners\nBolt frame", "#1565C0", "Day 1-2"),
        ("Phase 2\nY-Axis", "Install Y-motors\nRoute Y-belts\nAdd wheels", "#E53935", "Day 3-4"),
        ("Phase 3\nX-Axis", "Mount X-rail\nInstall carriage\nRoute X-belt", "#FF8F00", "Day 5-6"),
        ("Phase 4\nBalance", "Add pulleys\nHang weights\nOne-finger test", "#2E7D32", "Day 6-7"),
        ("Phase 5\nElectronics", "Flash GRBL\nWire motors\nWire switches", "#7B1FA2", "Day 7-8"),
        ("Phase 6\nCalibrate", "Home machine\nTune GRBL $\nTest square", "#00695C", "Day 9-10"),
    ]

    for i, (title, desc, color, days) in enumerate(phases):
        x = i * 2 + 0.5
        box = FancyBboxPatch((x, 0.3), 1.6, 2.2, boxstyle="round,pad=0.1",
                              facecolor=color, edgecolor="black", lw=1.5, alpha=0.9, zorder=3)
        ax.add_patch(box)
        ax.text(x + 0.8, 2.1, title, ha="center", va="center", fontsize=7,
                color="white", fontweight="bold", zorder=4)
        ax.text(x + 0.8, 1.2, desc, ha="center", va="center", fontsize=5.5,
                color="white", zorder=4)
        ax.text(x + 0.8, 0.5, days, ha="center", va="center", fontsize=6,
                color="#FFE082", fontweight="bold", zorder=4)

        if i < len(phases) - 1:
            ax.annotate("", xy=(x + 1.7, 1.4), xytext=(x + 1.55, 1.4),
                        arrowprops=dict(arrowstyle="->", color="#333", lw=2))

    plt.tight_layout()
    path = os.path.join(IMGDIR, "assembly_phases.png")
    fig.savefig(path, dpi=200, bbox_inches="tight")
    plt.close(fig)
    return path


def draw_roadmap():
    """Project phases / roadmap diagram."""
    fig, ax = plt.subplots(figsize=(10, 4.5))
    ax.set_xlim(0, 12)
    ax.set_ylim(-0.5, 5)
    ax.axis("off")
    ax.set_title("Project Roadmap", fontsize=14, fontweight="bold", pad=12)

    phases = [
        {
            "title": "Phase 1\n(CURRENT)",
            "subtitle": "B&W Sketching",
            "items": [
                "Single black pen",
                "Mechanical build",
                "Image-to-GCode pipeline",
                "Basic line drawings",
                "Calibration & tuning",
            ],
            "color": "#2E7D32",
            "border": "#1B5E20",
            "highlight": True,
        },
        {
            "title": "Phase 2\n(FUTURE)",
            "subtitle": "Multi-Color",
            "items": [
                "Pen dock mechanism",
                "Auto pen switching",
                "Color layer separation",
                "Registration accuracy",
                "Software color mapping",
            ],
            "color": "#1565C0",
            "border": "#0D47A1",
            "highlight": False,
        },
        {
            "title": "Phase 3\n(FUTURE)",
            "subtitle": "Advanced Output",
            "items": [
                "Fill patterns / hatching",
                "Halftoning (shading)",
                "Variable pen pressure",
                "Raspberry Pi sender",
                "Web UI / WiFi control",
            ],
            "color": "#7B1FA2",
            "border": "#4A148C",
            "highlight": False,
        },
    ]

    for i, phase in enumerate(phases):
        x = i * 4 + 0.3
        w = 3.4
        h = 4.0

        # Box
        alpha = 1.0 if phase["highlight"] else 0.7
        box = FancyBboxPatch((x, 0.3), w, h, boxstyle="round,pad=0.15",
                              facecolor=phase["color"], edgecolor=phase["border"],
                              lw=3 if phase["highlight"] else 1.5, alpha=alpha, zorder=3)
        ax.add_patch(box)

        # "CURRENT" banner
        if phase["highlight"]:
            banner = FancyBboxPatch((x + 0.2, 3.6), w - 0.4, 0.5, boxstyle="round,pad=0.05",
                                    facecolor="#FDD835", edgecolor="#F9A825", lw=1, zorder=4)
            ax.add_patch(banner)
            ax.text(x + w/2, 3.85, "THIS IS WHERE WE ARE", ha="center", va="center",
                    fontsize=6.5, color="#333", fontweight="bold", zorder=5)

        # Title
        ax.text(x + w/2, 3.2, phase["title"], ha="center", va="center",
                fontsize=9, color="white", fontweight="bold", zorder=4)

        # Subtitle
        ax.text(x + w/2, 2.55, phase["subtitle"], ha="center", va="center",
                fontsize=8, color="#FFE082", fontweight="bold", zorder=4)

        # Items
        for j, item in enumerate(phase["items"]):
            marker = ">" if phase["highlight"] else "-"
            ax.text(x + 0.3, 2.0 - j * 0.4, f"{marker} {item}",
                    fontsize=6.5, color="white", zorder=4, va="center")

        # Arrow to next phase
        if i < len(phases) - 1:
            ax.annotate("", xy=(x + w + 0.15, 2.3), xytext=(x + w + 0.05, 2.3),
                        arrowprops=dict(arrowstyle="->", color="#333", lw=2.5))

    plt.tight_layout()
    path = os.path.join(IMGDIR, "roadmap.png")
    fig.savefig(path, dpi=200, bbox_inches="tight")
    plt.close(fig)
    return path


def draw_software_pipeline():
    """Software pipeline diagram."""
    fig, ax = plt.subplots(figsize=(10, 2.5))
    ax.set_xlim(0, 14)
    ax.set_ylim(0, 3)
    ax.axis("off")
    ax.set_title("Software Pipeline: Image → Drawing", fontsize=12, fontweight="bold", pad=8)

    steps = [
        ("Input\nImage", "#333"),
        ("Grayscale\n+\nPreprocess", "#1565C0"),
        ("Edge\nDetect\n(XDoG)", "#7B1FA2"),
        ("Contour\nExtract", "#E53935"),
        ("Scale &\nClip to\nWork Area", "#FF8F00"),
        ("Path\nSort\n(nearest)", "#2E7D32"),
        ("G-Code\nGenerate", "#00695C"),
        ("Send to\nArduino\n(GRBL)", "#333"),
    ]

    for i, (label, color) in enumerate(steps):
        x = i * 1.7 + 0.3
        box = FancyBboxPatch((x, 0.5), 1.4, 1.8, boxstyle="round,pad=0.08",
                              facecolor=color, edgecolor="black", lw=1, alpha=0.9, zorder=3)
        ax.add_patch(box)
        ax.text(x + 0.7, 1.4, label, ha="center", va="center", fontsize=6,
                color="white", fontweight="bold", zorder=4)
        if i < len(steps) - 1:
            ax.annotate("", xy=(x + 1.5, 1.4), xytext=(x + 1.35, 1.4),
                        arrowprops=dict(arrowstyle="->", color="#333", lw=2))

    plt.tight_layout()
    path = os.path.join(IMGDIR, "software_pipeline.png")
    fig.savefig(path, dpi=200, bbox_inches="tight")
    plt.close(fig)
    return path


# ═══════════════════════════════════════════════════════════════
# DOCUMENT BUILDER
# ═══════════════════════════════════════════════════════════════

def set_cell_shading(cell, color_hex):
    """Set cell background color."""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), color_hex)
    shading.set(qn("w:val"), "clear")
    cell._tc.get_or_add_tcPr().append(shading)


def add_styled_table(doc, headers, rows, col_widths=None, header_color="1565C0"):
    """Add a formatted table."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"

    # Header row
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.size = Pt(9)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_shading(cell, header_color)

    # Data rows
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(str(val))
            run.font.size = Pt(8)
            if r_idx % 2 == 1:
                set_cell_shading(cell, "F5F5F5")

    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)

    doc.add_paragraph()  # spacer
    return table


def build_document():
    """Build the complete Word document."""
    print("Generating diagrams...")
    img_overview = draw_system_overview()
    img_arches = draw_architecture_comparison()
    img_cw = draw_direct_drive_diagram()
    img_pen = draw_pen_mechanism()
    img_elec = draw_electronics_diagram()
    img_phases = draw_assembly_phases()
    img_sw = draw_software_pipeline()
    img_roadmap = draw_roadmap()
    print("Diagrams done.")

    doc = Document()

    # ── Styles ──
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10)
    style.paragraph_format.space_after = Pt(4)

    for level in range(1, 4):
        hs = doc.styles[f"Heading {level}"]
        hs.font.color.rgb = RGBColor(0x15, 0x65, 0xC0)

    # ══════════════════════════════════════════════════════════
    # TITLE PAGE
    # ══════════════════════════════════════════════════════════
    for _ in range(6):
        doc.add_paragraph()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("ArtBot")
    run.bold = True
    run.font.size = Pt(36)
    run.font.color.rgb = RGBColor(0x15, 0x65, 0xC0)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("Large-Scale Vertical Drawing Machine")
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    doc.add_paragraph()

    desc = doc.add_paragraph()
    desc.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = desc.add_run("Hardware Plan & Architecture Document")
    run.font.size = Pt(14)

    doc.add_paragraph()
    doc.add_paragraph()

    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = info.add_run("IIT Delhi — Abu Dhabi\nMarch 2026")
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════
    # TABLE OF CONTENTS (manual)
    # ══════════════════════════════════════════════════════════
    doc.add_heading("Table of Contents", level=1)
    toc_items = [
        "1. Project Overview",
        "2. Project Roadmap",
        "3. System Architecture",
        "4. Architecture Comparison (Why Dual-Y Cartesian?)",
        "5. Drive System Comparison",
        "6. Motor Selection",
        "7. Stepper Driver Comparison",
        "8. Pen Lift Mechanism",
        "9. Controller & Firmware",
        "10. Rail & Frame Selection",
        "11. Direct-Drive Y-Axis (No Counterweights)",
        "12. Electronics & Wiring",
        "13. Software Pipeline",
        "14. Bill of Materials",
        "15. Assembly Plan",
        "16. GRBL Configuration",
        "17. Failure Modes & Mitigations",
        "18. References & Resources",
    ]
    for item in toc_items:
        p = doc.add_paragraph(item)
        p.paragraph_format.space_after = Pt(2)
        p.runs[0].font.size = Pt(10)

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════
    # 1. PROJECT OVERVIEW
    # ══════════════════════════════════════════════════════════
    doc.add_heading("1. Project Overview", level=1)

    doc.add_paragraph(
        "ArtBot is a large-scale robotic drawing machine that converts digital images into "
        "pen drawings on a large vertical surface (whiteboard, paper, canvas). Unlike typical CNC plotters that operate on "
        "horizontal surfaces, this machine works on a vertical plane — meaning the mechanical "
        "design must account for gravity acting on all moving components."
    )

    doc.add_heading("Key Specifications", level=2)
    add_styled_table(doc,
        ["Parameter", "Value"],
        [
            ["Board Size", "180 cm × 120 cm (vertical)"],
            ["Buffer Zone", "5 cm on all sides"],
            ["Usable Drawing Area", "170 cm × 110 cm"],
            ["Origin", "Top-left corner"],
            ["Motion System", "Dual-Y Cartesian (no counterweights)"],
            ["Motors", "3× NEMA 23 steppers (no counterweights needed) + 1× SG90 servo"],
            ["Controller", "Arduino Uno + CNC Shield V3 + GRBL firmware"],
            ["Drivers", "3× TB6600"],
            ["Resolution", "80 steps/mm (0.0125 mm per step)"],
            ["Pen Tip Width", "0.5–1.0 mm (40–80× larger than step size)"],
        ],
        col_widths=[5, 10],
    )

    doc.add_heading("The Three Axes", level=2)
    add_styled_table(doc,
        ["Axis", "Function", "Mechanism"],
        [
            ["X", "Horizontal (left–right)", "Belt + motor on rail"],
            ["Y", "Vertical (up–down)", "Dual NEMA 23 motors + belts (direct drive)"],
            ["Z", "Pen up/down", "SG90 servo + compression spring"],
        ],
        col_widths=[2, 5, 8],
    )

    # ══════════════════════════════════════════════════════════
    # 2. PROJECT ROADMAP
    # ══════════════════════════════════════════════════════════
    doc.add_heading("2. Project Roadmap", level=1)
    doc.add_paragraph(
        "ArtBot is being developed in phases. The current focus is Phase 1: getting the machine "
        "to produce clean black-and-white line drawings with a single pen. Multi-color pen switching "
        "and advanced output techniques are planned for future phases."
    )
    doc.add_paragraph(
        "Our architecture choices (Dual-Y Cartesian, rigid rails, precise homing) were made with "
        "all three phases in mind — so we won't have to rebuild the machine when we add capabilities."
    )
    doc.add_picture(img_roadmap, width=Inches(6.0))
    last_paragraph = doc.paragraphs[-1]
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_heading("Phase 1 — Black & White Sketching (Current)", level=2)
    for item in [
        "Single black pen (marker, felt-tip, or ballpoint)",
        "Mechanical skeleton: frame, rails, NEMA 23 motors (no counterweights)",
        "Electronics: Arduino + CNC Shield + GRBL + TB6600 drivers",
        "Software pipeline: image to edge detection to contours to G-Code",
        "Calibration: homing, steps/mm tuning, test square verification",
        "Goal: produce clean, recognizable line drawings on the board",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_heading("Phase 2 — Multi-Color (Future)", level=2)
    for item in [
        "Pen dock mechanism at top of board (row of pen holders at known coordinates)",
        "Automatic pen switching: carriage moves to dock, drops current pen, picks up next color",
        "Color layer separation in software (split image into color channels → separate G-Code passes)",
        "Registration accuracy: machine must return to exact coordinates between pen changes",
        "This is why we chose rigid rails over cables — cables can't guarantee return precision",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_heading("Phase 3 — Advanced Output (Future)", level=2)
    for item in [
        "Fill patterns and hatching (parallel lines at varying density for shading)",
        "Halftoning (dots of varying size/spacing to simulate grayscale)",
        "Variable pen pressure via Z-axis stepper upgrade (replaces servo)",
        "Raspberry Pi as wireless G-Code sender (web UI, file storage, WiFi control)",
        "Integration with vpype (plotter-specific path optimization library)",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    p = doc.add_paragraph()
    run = p.add_run(
        "Everything in this document focuses on Phase 1. "
        "Phases 2 and 3 are mentioned only to explain why certain architectural decisions "
        "were made (e.g., choosing rigid rails over cables for future pen-switch registration)."
    )
    run.italic = True

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════
    # 3. SYSTEM ARCHITECTURE
    # ══════════════════════════════════════════════════════════
    doc.add_heading("3. System Architecture", level=1)
    doc.add_paragraph(
        "The diagram below shows the complete machine layout as seen from the front. "
        "Two Y-motors at the top corners move the horizontal X-rail up and down. "
        "A third motor on the X-rail moves the pen carriage left and right. "
        "NEMA 23 motors drive the Y-axis directly against gravity (no counterweights needed)."
    )
    doc.add_picture(img_overview, width=Inches(5.5))
    last_paragraph = doc.paragraphs[-1]
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════
    # 3. ARCHITECTURE COMPARISON
    # ══════════════════════════════════════════════════════════
    doc.add_heading("4. Architecture Comparison", level=1)
    doc.add_paragraph(
        "We evaluated four motion architectures for a 180×120 cm vertical plotter. "
        "The diagram below shows all four side by side, followed by a detailed comparison."
    )
    doc.add_picture(img_arches, width=Inches(6.0))
    last_paragraph = doc.paragraphs[-1]
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_heading("A: Polargraph / V-Plotter — REJECTED", level=2)
    doc.add_paragraph(
        "Two motors at top corners suspend a pen gondola on cables. "
        "Used by: Makelangelo, Maslow CNC, Penelope, Stringent."
    )
    doc.add_paragraph("Why rejected:", style="List Bullet")
    for reason in [
        "Accuracy collapses at bottom corners — the Jacobian of the coordinate transform shows resolution drops dramatically below the motor line. Most builders recommend 15% margins, wasting ~30% of our board.",
        "Gondola swings on fast direction changes — must draw very slowly.",
        "Cable stretch accumulates over long drawings, causing drift.",
        "Multi-color pen switching is nearly impossible (gondola sways during docking).",
        "Penelope's builder reported accuracy issues 'especially near the bottom and on curves.'"
    ]:
        doc.add_paragraph(reason, style="List Bullet 2")

    doc.add_heading("B: CoreXY / H-Bot — REJECTED", level=2)
    doc.add_paragraph(
        "Two stationary motors drive a single continuous belt loop. "
        "Used by: most modern 3D printers (Voron, Bambu Lab), Stefan van Seggelen's large plotter."
    )
    for reason in [
        "At 180×120 cm, total belt path is ~12 meters — belt stretch is unmanageable with standard GT2.",
        "Stefan van Seggelen's build required 15mm steel-core GT2 belts (expensive, hard to source).",
        "Getting two 6-meter belt paths tensioned identically is nearly impossible for a first build.",
        "H-Bot variant introduces rotational torque on diagonal moves, causing skew."
    ]:
        doc.add_paragraph(reason, style="List Bullet 2")

    doc.add_heading("C: Single-Y Cartesian — REJECTED", level=2)
    doc.add_paragraph(
        "One motor per axis — simplest possible setup."
    )
    for reason in [
        "Single motor driving one side of a 180 cm rail causes severe racking (rail tilts).",
        "No large-format CNC machine uses single-Y drive — it simply doesn't work at this scale."
    ]:
        doc.add_paragraph(reason, style="List Bullet 2")

    doc.add_heading("D: Dual-Y Cartesian — CHOSEN", level=2)
    doc.add_paragraph(
        "Two synchronized Y-motors (one per side) move the X-rail vertically. "
        "A third motor drives the pen carriage horizontally. NEMA 23 motors handle gravity directly. "
        "Used by: iDraw H (commercial, ±0.02mm precision), OpenBuilds ACRO, every large CNC router."
    )
    for reason in [
        "Anti-racking: both sides of the rail move simultaneously, keeping it perfectly level.",
        "Short belt runs (~260 cm each) — minimal stretch compared to CoreXY's 12m.",
        "No counterweights needed: NEMA 23 motors provide sufficient torque against gravity.",
        "Pen-switch ready: rigid rail means precise return to coordinates for multi-color.",
        "GRBL natively supports dual-Y axis with the CNC Shield's A-axis clone jumper.",
        "Proven at scale by commercial and DIY machines worldwide."
    ]:
        doc.add_paragraph(reason, style="List Bullet 2")

    doc.add_page_break()

    # Summary table
    doc.add_heading("Architecture Decision Matrix", level=2)
    add_styled_table(doc,
        ["Criteria", "Polargraph", "CoreXY", "Single-Y", "Dual-Y (Ours)"],
        [
            ["Precision at corners", "Poor", "Good", "Poor", "Excellent"],
            ["Scale to 180cm", "Easy", "Hard", "OK", "Proven"],
            ["Belt stretch risk", "N/A (cables)", "Severe", "Low", "Low"],
            ["Anti-racking", "N/A", "Good", "None", "Excellent"],
            ["Multi-color future", "Very hard", "Possible", "Hard", "Easy"],
            ["Build complexity", "Low", "High", "Low", "Medium"],
            ["Cost", "$15–100", "$200+", "$100", "$150–200"],
            ["Community support", "Large", "Medium", "N/A", "Large"],
        ],
        col_widths=[3.5, 2.5, 2.5, 2.5, 3],
    )

    # ══════════════════════════════════════════════════════════
    # 4. DRIVE SYSTEM
    # ══════════════════════════════════════════════════════════
    doc.add_heading("5. Drive System Comparison", level=1)
    add_styled_table(doc,
        ["Drive Type", "Speed", "Stretch", "Self-Lock (Gravity Safe)", "Cost", "Verdict"],
        [
            ["GT2 Timing Belt", "Fast (5000 mm/min)", "~0.2%", "No — motor holding torque only", "$5–10 / 10m", "CHOSEN"],
            ["Lead Screw (T8)", "Slow (500 mm/min)", "~0", "Yes — self-locking", "$20–30 / 120cm", "Rejected (too slow)"],
            ["Ball Screw", "Fast", "~0", "No", "$100–300 / 120cm", "Rejected (too expensive)"],
            ["Cable / Fishing Line", "Medium", "High", "No", "$1", "Rejected (stretch)"],
            ["Beaded Chain", "Medium", "Low", "No", "$5–10", "Rejected (backlash)"],
            ["Rack & Pinion", "Fast", "~0", "No", "$50+ / 180cm", "Rejected (cost)"],
        ],
        col_widths=[3, 2.5, 1.5, 3, 2.5, 2.5],
    )
    doc.add_paragraph(
        "GT2 timing belt wins on speed, cost, and community support. "
        "The self-locking disadvantage is mitigated by NEMA 23 holding torque — "
        "the rail stays in place even with power off."
    )

    # ══════════════════════════════════════════════════════════
    # 5. MOTOR SELECTION
    # ══════════════════════════════════════════════════════════
    doc.add_heading("6. Motor Selection", level=1)
    add_styled_table(doc,
        ["Motor", "Frame Size", "Torque", "Weight", "Current", "Voltage", "Cost", "Verdict"],
        [
            ["NEMA 14", "35mm", "10–20 Ncm", "~150g", "0.5–1A", "12V", "$5–8", "Too weak"],
            ["NEMA 23", "57mm", "100–300 Ncm", "~600-1000g", "2.0–4A", "24V", "$20–40", "CHOSEN"],
            ["NEMA 23", "57mm", "100–300 Ncm", "~600g", "2–4A", "24V", "$20–40", "Overkill"],
            ["Closed-Loop Servo", "Varies", "High", "~400g", "Varies", "24V+", "$50+", "Too expensive"],
        ],
        col_widths=[2.5, 1.5, 2, 1.5, 1.5, 1.2, 1.5, 2.5],
    )
    doc.add_paragraph(
        "NEMA 23 (>100 Ncm) provides enough torque to drive the Y-axis directly against gravity, "
        "torque to position the rail. NEMA 23 would add unnecessary weight to the carriage and require "
        "24V power and higher-current drivers."
    )

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════
    # 6. STEPPER DRIVERS
    # ══════════════════════════════════════════════════════════
    doc.add_heading("7. Stepper Driver Comparison", level=1)
    add_styled_table(doc,
        ["Driver", "Max Current", "Microstepping", "Noise", "Cost", "Verdict"],
        [
            ["A4988", "2A (1A safe)", "1/16", "Very loud", "$1–2", "Rejected (noise)"],
            ["DRV8825", "2.2A (1.5A safe)", "1/32", "Loud", "$2–3", "Rejected (noise)"],
            ["TB6600", "4.0A", "1/32", "Moderate", "$8–12", "CHOSEN"],
        ],
        col_widths=[2.5, 2.5, 2, 2.5, 2, 3],
    )
    doc.add_paragraph(
        "TB6600 supports the higher current draw of NEMA 23 motors (up to 4A per phase) — required "
        "that may run 30+ minutes per drawing. The extra $1–2 per driver is a trivial cost. "
        "Drop-in compatible with CNC Shield V3 (same pinout as A4988)."
    )

    # ══════════════════════════════════════════════════════════
    # 7. PEN LIFT
    # ══════════════════════════════════════════════════════════
    doc.add_heading("8. Pen Lift Mechanism", level=1)
    doc.add_paragraph(
        "On a vertical board, gravity pulls the pen AWAY from the surface. "
        "A compression spring pushes the pen toward the board. "
        "The servo retracts the pen when traveling between strokes."
    )
    doc.add_picture(img_pen, width=Inches(5.5))
    last_paragraph = doc.paragraphs[-1]
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    add_styled_table(doc,
        ["Mechanism", "Weight", "Speed", "Power", "Cost", "Verdict"],
        [
            ["SG90 Servo", "9g", "0.1s / 60°", "5V (from Arduino)", "$1–2", "CHOSEN"],
            ["Solenoid", "30–80g", "5–10ms", "12V continuous", "$5–10", "Rejected (thumps board)"],
            ["Stepper Motor", "200g+", "Variable", "Needs extra driver", "$10+", "Rejected (too heavy)"],
        ],
        col_widths=[2.5, 1.5, 2, 3, 1.5, 3.5],
    )

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════
    # 8. CONTROLLER
    # ══════════════════════════════════════════════════════════
    doc.add_heading("9. Controller & Firmware", level=1)
    add_styled_table(doc,
        ["Option", "Real-Time?", "Dual-Y Support", "WiFi", "Complexity", "Verdict"],
        [
            ["Arduino Uno + GRBL", "Yes", "Native (A-axis clone)", "No", "Low", "CHOSEN (v1)"],
            ["Arduino Mega + Marlin", "Yes", "Yes", "No", "Medium", "Overkill (3D printer firmware)"],
            ["ESP32 + FluidNC", "Yes", "Yes", "Yes", "Medium", "Future upgrade path"],
            ["Raspberry Pi (alone)", "No", "N/A", "Yes", "High", "Rejected (not real-time)"],
        ],
        col_widths=[3, 1.5, 3, 1, 2, 3.5],
    )
    doc.add_paragraph(
        "GRBL on Arduino Uno is the proven standard for 3-axis CNC. It generates motor pulses "
        "with microsecond precision and has a look-ahead algorithm that smooths curves automatically. "
        "A Raspberry Pi can be added later as a wireless G-Code sender, but cannot replace the Arduino "
        "for real-time motor control."
    )

    # ══════════════════════════════════════════════════════════
    # 9. RAIL SELECTION
    # ══════════════════════════════════════════════════════════
    doc.add_heading("10. Rail & Frame Selection", level=1)
    add_styled_table(doc,
        ["Profile", "Cross Section", "Stiffness (Iy)", "Weight/m", "Deflection at 180cm, 1kg", "Use"],
        [
            ["V-Slot 2020", "20×20 mm", "~0.7 cm⁴", "~0.5 kg", "~4 mm (UNACCEPTABLE)", "Y-rails, frame"],
            ["V-Slot 2040", "20×40 mm", "~3.3 cm⁴", "~0.9 kg", "~0.9 mm (acceptable)", "X-RAIL (chosen)"],
            ["V-Slot 4040", "40×40 mm", "~7.0 cm⁴", "~1.6 kg", "~0.4 mm (excellent)", "Not needed"],
        ],
        col_widths=[2, 2, 2, 1.5, 3.5, 3],
    )
    doc.add_paragraph(
        "The X-rail (180 cm horizontal span) must be 2040 or stiffer. A 2020 rail will visibly bow "
        "at the center. Adding a flat aluminum backing strip (3mm × 30mm × 180cm) to the 2040 "
        "provides additional rigidity at almost no cost."
    )

    # ══════════════════════════════════════════════════════════
    # 10. COUNTERWEIGHT
    # ══════════════════════════════════════════════════════════
    doc.add_heading("11. Direct-Drive Y-Axis (No Counterweights)", level=1)
    doc.add_picture(img_cw, width=Inches(4.5))
    last_paragraph = doc.paragraphs[-1]
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_heading("Balancing Procedure", level=2)
    steps = [
        "Fully assemble the X-rail: rail + carriage + X-motor + servo + pen holder + all brackets.",
        "Weigh the complete assembly on a kitchen scale.",
        "No counterweights needed — NEMA 23 motors handle gravity directly.",
        "Verify motors hold X-rail position when powered (should not drift).",
        "The One-Finger Test: push the rail with one finger. It should move smoothly and stay exactly where you leave it.",
        "If the rail drifts down: add weight. If it drifts up: remove weight.",
        "Note: X-rail will drop on power loss — acceptable for pen plotter.",
    ]
    for i, step in enumerate(steps, 1):
        doc.add_paragraph(f"{i}. {step}")

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════
    # 11. ELECTRONICS
    # ══════════════════════════════════════════════════════════
    doc.add_heading("12. Electronics & Wiring", level=1)
    doc.add_picture(img_elec, width=Inches(5.5))
    last_paragraph = doc.paragraphs[-1]
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_heading("Critical Wiring Note — Dual Y-Motors", level=2)
    p = doc.add_paragraph()
    run = p.add_run(
        "WARNING: The two Y-motors face opposite directions physically. If you clone them with "
        "identical signals (via the A-axis jumper), one will push UP while the other pushes DOWN. "
        "The rail will jam and potentially break."
    )
    run.bold = True
    run.font.color.rgb = RGBColor(0xC6, 0x28, 0x28)

    doc.add_paragraph(
        "Fix: On ONE of the two Y-motors, swap the wiring of one coil pair. If the motor connector "
        "has 4 wires [A+, A−, B+, B−], swap A+ and A− on that one motor. This reverses its rotation, "
        "making both motors move the rail in the same direction."
    )

    # ══════════════════════════════════════════════════════════
    # 12. SOFTWARE PIPELINE
    # ══════════════════════════════════════════════════════════
    doc.add_heading("13. Software Pipeline", level=1)
    doc.add_paragraph(
        "The software converts a digital image into G-Code instructions that the Arduino executes."
    )
    doc.add_picture(img_sw, width=Inches(6.0))
    last_paragraph = doc.paragraphs[-1]
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph(
        "The pipeline is implemented in Python and is fully functional. "
        "Three CLI tools are provided:"
    )
    for tool in [
        "convert.py — converts an image to .gcode (the main pipeline)",
        "simulate.py — visualizes the G-Code as an animated pen plotter preview",
        "tune.py — interactive slider interface for tuning edge detection parameters",
    ]:
        doc.add_paragraph(tool, style="List Bullet")

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════
    # 13. BOM
    # ══════════════════════════════════════════════════════════
    doc.add_heading("14. Bill of Materials", level=1)

    doc.add_heading("Structural", level=2)
    add_styled_table(doc,
        ["Item", "Spec", "Qty", "Purpose"],
        [
            ["V-Slot 2040 Extrusion", "180 cm", "1", "X-axis horizontal rail"],
            ["V-Slot 2020 Extrusion", "120 cm", "2", "Y-axis vertical rails"],
            ["V-Slot 2020 Extrusion", "180 cm", "2", "Frame top & bottom"],
            ["90° Angle Brackets", "V-Slot compatible", "8", "Corner joints"],
            ["T-Nuts + M5×8 Bolts", "M5", "50+ pack", "Securing components"],
            ["Flat Aluminum Bar", "3mm × 30mm × 180cm", "1", "X-rail stiffener"],
        ],
        col_widths=[4, 3, 1, 5],
    )

    doc.add_heading("Motion", level=2)
    add_styled_table(doc,
        ["Item", "Spec", "Qty", "Purpose"],
        [
            ["GT2 Timing Belt", "6mm wide, open-ended", "10m roll", "All axis drive"],
            ["GT2 Pulley (20-tooth)", "5mm bore", "3", "Motor shafts"],
            ["GT2 Idler Pulley", "5mm bore", "5", "Belt routing"],
            ["V-Slot Gantry Plate", "Standard OpenBuilds", "3", "Rail carriages"],
            ["Mini V-Wheels (Delrin)", "Standard", "12", "4 per gantry plate"],
            ["Eccentric Spacers", "Standard", "6", "Wheel tension adjust"],
            ["Belt Tensioners", "Adjustable, spring-loaded", "3", "One per belt run"],
        ],
        col_widths=[4, 3, 1.5, 4.5],
    )

    doc.add_heading("Motors & Electronics", level=2)
    add_styled_table(doc,
        ["Item", "Spec", "Qty", "Purpose"],
        [
            ["NEMA 23 Stepper Motor", ">100 Ncm, 2.0–4A", "3", "X + Y-Left + Y-Right"],
            ["SG90 Micro Servo", "Standard", "1", "Pen lift (Z-axis)"],
            ["Motor Mounting Bracket", "NEMA 23 L-bracket (57mm)", "3", "Motor mounting"],
            ["Arduino Uno R3", "Original or clone", "1", "GRBL controller"],
            ["CNC Shield V3", "Fits Arduino Uno", "1", "Driver interface"],
            ["TB6600 Driver", "External, up to 4A", "3", "Motor drivers"],
            ["24V 15A Power Supply", "Mean Well or equiv.", "1", "Main power"],
            ["Limit Switches", "NO microswitch, lever", "3", "Homing (X, Y, spare)"],
            ["Drag Chain", "10×10mm flexible", "2m", "Cable management"],
            ["USB-B Cable", "Standard Arduino", "1", "PC connection"],
        ],
        col_widths=[4, 3, 1, 5],
    )

    doc.add_heading("Pen Mechanism", level=2)
    add_styled_table(doc,
        ["Item", "Spec", "Qty", "Purpose"],
        [
            ["Compression Spring", "~10mm OD, ~20mm", "1", "Push pen toward board"],
            ["Pen Holder", "3D-printed or purchased", "1", "Holds marker / pen"],
        ],
        col_widths=[4, 3, 1.5, 4.5],
    )
    doc.add_paragraph("Note: No counterweight system needed — NEMA 23 motors handle gravity directly.")

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════
    # 14. ASSEMBLY PLAN
    # ══════════════════════════════════════════════════════════
    doc.add_heading("15. Assembly Plan", level=1)
    doc.add_picture(img_phases, width=Inches(6.0))
    last_paragraph = doc.paragraphs[-1]
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    phases = [
        ("Phase 1: Frame (Day 1–2)", [
            "Lay out all 4 frame extrusions on the floor.",
            "Bolt corners with angle brackets + T-nuts.",
            "CRITICAL: Measure BOTH diagonals — they must be equal (within 2mm) or drawings will be skewed.",
            "Mount frame to drawing surface or backing structure with at least 6 mounting points.",
        ]),
        ("Phase 2: Y-Axis (Day 3–4)", [
            "Install V-wheels on the two Y-axis gantry plates.",
            "Slide gantry plates onto left and right vertical rails.",
            "Adjust eccentric spacers: wheels should grip snugly but roll freely.",
            "Mount Y-Left motor at top-left, Y-Right motor at top-right.",
            "Install idler pulleys at the bottom of each vertical rail.",
            "Route Y-axis belts: motor → down to idler → back up to gantry plate.",
            "Test: turn both motor shafts BY HAND — both gantry plates should move smoothly.",
        ]),
        ("Phase 3: X-Axis (Day 5–6)", [
            "Bolt the 180cm 2040 rail between the two Y-axis gantry plates.",
            "Check level — the X-rail must be perfectly horizontal.",
            "Install pen carriage on the X-rail with V-wheels.",
            "Mount X-motor at the left end, idler pulley at the right end.",
            "Route X-axis belt and install belt tensioner.",
            "Optional: bolt mid-span aluminum stiffener to back of X-rail.",
        ]),
        ("Phase 4: Electronics (Day 6–7)", [
            "Flash GRBL firmware onto the Arduino Uno.",
            "Plug CNC Shield onto Arduino, wire TB6600 external drivers.",
            "Set A-axis jumper on CNC Shield to clone Y.",
            "Wire motors: X→X slot, Y-Left→Y slot, Y-Right→A slot.",
            "REVERSE ONE Y-MOTOR'S COIL (swap A+ and A− on one motor).",
            "Wire servo to SpnEn pin.",
            "Wire and mount limit switches (X-min, Y-min).",
            "Connect 24V 15A PSU to TB6600 drivers.",
        ]),
        ("Phase 5: Calibration (Day 8–9)", [
            "Connect to GRBL via Universal Gcode Sender.",
            "Configure GRBL $ parameters (see Section 15).",
            "Home the machine ($H) — should move to top-left and stop at switches.",
            "Test X and Y directions with small jog commands.",
            "Draw a test square and measure with a ruler.",
            "If dimensions are wrong, adjust $100 / $101 (steps/mm).",
        ]),
    ]

    for phase_title, steps in phases:
        doc.add_heading(phase_title, level=2)
        for step in steps:
            p = doc.add_paragraph(step, style="List Bullet")
            if "CRITICAL" in step or "REVERSE" in step:
                for run in p.runs:
                    run.bold = True
                    run.font.color.rgb = RGBColor(0xC6, 0x28, 0x28)

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════
    # 15. GRBL CONFIG
    # ══════════════════════════════════════════════════════════
    doc.add_heading("16. GRBL Configuration", level=1)
    doc.add_paragraph("Enter these via the serial console in Universal Gcode Sender:")
    add_styled_table(doc,
        ["Parameter", "Value", "Description"],
        [
            ["$100", "80", "X steps/mm (200 steps × 16 microsteps / 20 teeth × 2mm)"],
            ["$101", "80", "Y steps/mm (same calculation)"],
            ["$110", "5000", "X max feed rate (mm/min)"],
            ["$111", "5000", "Y max feed rate (mm/min)"],
            ["$120", "200", "X acceleration (mm/s²) — start low, increase gradually"],
            ["$121", "100", "Y acceleration (mm/s²) — lower than X (heavy rail)"],
            ["$130", "1700", "X max travel (mm)"],
            ["$131", "1100", "Y max travel (mm)"],
            ["$22", "1", "Enable homing cycle"],
            ["$23", "3", "Homing direction invert (both X and Y home to min)"],
            ["$24", "25", "Homing feed rate (mm/min) — slow approach to switch"],
            ["$25", "500", "Homing seek rate (mm/min) — fast initial move"],
        ],
        col_widths=[2, 1.5, 10],
    )

    doc.add_heading("Test Square G-Code", level=2)
    doc.add_paragraph("Send this to verify calibration. Draws a 150mm square — measure with a ruler:")
    test_code = (
        "G21           ; Units: mm\n"
        "G90           ; Absolute positioning\n"
        "G0 X50 Y50 F5000    ; Move to start\n"
        "M3 S90        ; Pen down\n"
        "G4 P0.15      ; Wait for pen settle\n"
        "G1 X200 Y50 F3000   ; Draw right\n"
        "G1 X200 Y200        ; Draw down\n"
        "G1 X50 Y200         ; Draw left\n"
        "G1 X50 Y50          ; Draw up (close square)\n"
        "M5            ; Pen up\n"
        "G0 X0 Y0 F5000      ; Return home\n"
        "M2            ; Program end"
    )
    p = doc.add_paragraph()
    run = p.add_run(test_code)
    run.font.name = "Consolas"
    run.font.size = Pt(8)

    # ══════════════════════════════════════════════════════════
    # 16. FAILURE MODES
    # ══════════════════════════════════════════════════════════
    doc.add_heading("17. Failure Modes & Mitigations", level=1)
    add_styled_table(doc,
        ["Failure", "Cause", "Prevention"],
        [
            ["Rail drops on power loss", "Motors lose holding torque", "Acceptable for pen plotter — keep motors enabled during jobs"],
            ["Skipped steps (drawing shifts)", "Accel too high, PSU weak, belt loose", "Tune $120/$121, use 10A PSU, tension belts"],
            ["Y-axis jams on startup", "Y-motors fighting each other", "Reverse one motor's coil wiring"],
            ["Drawing is skewed (rhombus)", "Frame not square", "Measure diagonals during assembly"],
            ["Pen doesn't touch board", "Spring too weak / servo miscalibrated", "Tune spring, calibrate servo angle"],
            ["Belt skips teeth", "Belt loose, pulley grub screw loose", "Belt tensioners + Loctite on grub screws"],
            ["X-rail bows at center", "180cm 2020 profile too flexible", "Use 2040 + mid-span stiffener"],
            ["Lines are wavy/vibrating", "Rail chatter at high speed", "Reduce feed rate, stiffen rail, lower accel"],
            ["Homing fails", "Switch not triggering / wrong config", "Test switches manually, check $22/$23"],
        ],
        col_widths=[3, 3.5, 5.5],
    )

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════
    # 17. REFERENCES
    # ══════════════════════════════════════════════════════════
    doc.add_heading("18. References & Resources", level=1)

    doc.add_heading("Real-World Builds & Projects", level=2)
    refs = [
        ("Makelangelo (Polargraph)", "github.com/MarginallyClever/Makelangelo-software"),
        ("Maslow CNC (Wall-Mounted Router)", "maslowcnc.com"),
        ("Penelope Vertical Pen Plotter", "davidbliss.com/2021/09/13/penelope"),
        ("Stefan van Seggelen Large CoreXY Plotter", "stefanvanseggelen.com/plotter.html"),
        ("iDraw H Large Format Plotter", "uunatek.com/collections/large-format-size-drawing-robots"),
        ("OpenBuilds Pen Plotter", "builds.openbuilds.com/builds/pen-plotter.10615"),
        ("OpenBuilds ACRO System", "builds.openbuilds.com/tags/acro"),
        ("Jason Webb's GRBL Wall Plotter", "github.com/jasonwebb/grbl-mega-wall-plotter"),
        ("Stringent $15 Wall Plotter", "hackster.io/fredrikstridsman/stringent-the-15-wall-plotter-d965ca"),
    ]
    for name, url in refs:
        doc.add_paragraph(f"{name} — {url}", style="List Bullet")

    doc.add_heading("Technical References", level=2)
    tech_refs = [
        ("GRBL Setup Guide", "howtomechatronics.com/tutorials/how-to-setup-grbl-control-cnc-machine-with-arduino"),
        ("GRBL with CNC Shield", "diyengineers.com/2023/01/05/grbl-with-arduino-cnc-shield-complete-guide"),
        ("CoreXY Belt Tensioning", "drmrehorst.blogspot.com/2018/08/corexy-mechanism-layout-and-belt.html"),
        ("Belts vs Leadscrews in CNC", "blanch.org/belts-vs-screws-in-cnc-design"),
        ("V-Slot Deflection Calculator", "builds.openbuilds.com/threads/how-to-calculate-v-slot-deflection.4881"),
        ("NEMA 23 vs NEMA 23", "sss-motors.com/nema-17-vs-nema-23-stepper-motor"),
        ("TB6600 vs TMC2208", "xecor.com/blog/tmc2208-vs-tmc2209"),
        ("Polargraph Mathematics", "nexp.pt/vbot.html"),
        ("FluidNC (Modern GRBL Replacement)", "wiki.fluidnc.com/en/config/kinematics"),
    ]
    for name, url in tech_refs:
        doc.add_paragraph(f"{name} — {url}", style="List Bullet")

    # ── Save ──
    outpath = os.path.join(OUTDIR, "ArtBot_Hardware_Plan.docx")
    doc.save(outpath)
    print(f"\nDocument saved: {outpath}")
    return outpath


if __name__ == "__main__":
    build_document()
